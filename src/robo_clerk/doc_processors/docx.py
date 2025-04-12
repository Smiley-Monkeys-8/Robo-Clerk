import os
import re
from typing import List
from docx import Document
import unicodedata

from robo_clerk.doc_processors.types import Feature

class DOCXProcessor:
    def __init__(self, file_path):
        self.file_path = file_path
        self.text = ""
        self.cleaned_text = ""
        self.extracted_info = {}
        self.tickable_info = {}

    # 1. Extract raw text
    def extract_text(self):
        self.text = self.extract_docx_text_with_checkboxes(self.file_path)

    # 2. Clean the extracted text
    def clean_text(self):
        text = unicodedata.normalize("NFKD", self.text)
        text = re.sub(r"[^a-zA-Z0-9☐☒.,!?%€$-:\n\s@()]", " ", text)  # preserve checkboxes and special chars
        text = text.lower()
        text = re.sub(r"\s+", " ", text).strip()
        self.cleaned_text = text

    # 3. Extract text input
    def extract_info(self):
        self.personal_info={}
        personal_info_patterns = {
            "last_name": r"last name\s*([\w\-]+)(?=\s*(first/ middle name \(s\)|address|country of domicile|date of birth|nationality|passport no|id type|id issue date|id expiry date))",
            "first_name": r"first/ middle name \(s\)\s*([\w\s]+)(?=\s*(address|country of domicile|date of birth|nationality|passport no|id type|id issue date|id expiry date))",
            "address": r"address\s*([\w\s,.-]+?)(?=\s*(country of domicile|date of birth|nationality|passport no|id type|id issue date|id expiry date))",
            "country_of_domicile": r"country of domicile\s*([\w\s]+)(?=\s*(date of birth|nationality|passport no|id type|id issue date|id expiry date))",
            "date_of_birth": r"date of birth\s*(\d{4}-\d{2}-\d{2})(?=\s*(nationality|passport no|id type|id issue date|id expiry date))",
            "nationality": r"nationality\s*([\w\s]+)(?=\s*(passport no|id type|id issue date|id expiry date))",
            "passport_no": r"passport no/ unique id\s*([\w\d]+)(?=\s*(id type|id issue date|id expiry date))",
            "id_type": r"id type\s*([\w\s]+)(?=\s*(id issue date|id expiry date))",
            "id_issue_date": r"id issue date\s*(\d{4}-\d{2}-\d{2})(?=\s*id expiry date)",
            "id_expiry_date": r"id expiry date\s*(\d{4}-\d{2}-\d{2})",
            "telephone": r"telephone\s*(\+[\d\s]+)(?=\s*(e-mail|email))",
            "email": r"(e[-\s]*mail|email)\s*([\w\.-]+@[\w\.-]+)(?=\s*(telephone|country of domicile|date of birth|nationality|passport no|id type))",
        }

        # Extract personal information
        for key, pattern in personal_info_patterns.items():
            match = re.search(pattern, self.cleaned_text)
            if match:
                if key == "first_name":
                    full_name = match.group(1).strip()
                    names = full_name.split()
                    self.personal_info["first_name"] = names[0]
                    self.personal_info["middle_name"] = " ".join(names[1:]) if len(names) > 1 else ""
                else:
                    self.personal_info[key] = match.group(1)

        # Extract tickable fields
        self.tickable_info = self.extract_checked_options_from_text(self.text)

    # Extract communication fields (telephone, email)
    def extract_communication_info(self):
        self.communication_info = {}
        communication_patterns = {
            "telephone": r"telephone\s*([\+0-9\s\-\(\)]+)(?=\s*(e[-\s]*mail|email))",
            "email": r"(e[-\s]*mail|email)\s*([\w\.-]+@[\w\.-]+)"
        }

        # Extract communication information
        for key, pattern in communication_patterns.items():
            match = re.search(pattern, self.cleaned_text)
            if match:
                self.communication_info[key] = match.group(2).strip()  # Adjusted to capture the email address correctly

    # Methods to extract the text options from checkboxes
    def extract_docx_text_with_checkboxes(self, file_path):
        try:
            doc = Document(file_path)
        except Exception as e:
            print(f"Error reading DOCX file: {e}")
            return ""

        all_text = []

        for i, table in enumerate(doc.tables):
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_runs = []
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            cell_runs.append(run.text)
                    full_cell_text = "".join(cell_runs).strip()
                    if full_cell_text:
                        row_text.append(full_cell_text)
                if row_text:
                    all_text.append(" | ".join(row_text))

        return "\n".join(all_text)

    # Function to extract the checked options from the text
    def extract_checked_options_from_text(self, text):
        checkbox_info = {}

        lines = text.split("\n")
        for line in lines:
            if "☒" in line:
                parts = line.split("|")
                label = parts[0].strip().lower() if len(parts) > 1 else "unknown"
                options_text = " | ".join(parts[1:]) if len(parts) > 1 else parts[0]

                matches = re.findall(r"(☒|☐)\s*([^\t|☐☒]+)", options_text)

                checked = [opt.strip().lower() for mark, opt in matches if mark == "☒"]

                if checked:
                    checkbox_info[label] = checked if len(checked) > 1 else checked[0]

        return checkbox_info
        

    # Run all steps
    def run_pipeline(self) -> List[Feature]:
        self.extract_text()
        self.clean_text()
        self.extract_text()
        self.clean_text()
        self.extract_info()
        self.extract_communication_info()
        
        final_info = {**self.personal_info, **self.tickable_info, **self.communication_info}

        features: List[Feature] = [Feature(key=key, value=value, source=self.file_path) for key, value in final_info.items()]

        return features
